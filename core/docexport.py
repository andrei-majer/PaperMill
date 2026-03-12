"""Markdown → .docx export with academic formatting."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import (
    DOCX_FONT, DOCX_FONT_SIZE_PT, DOCX_LINE_SPACING, DOCX_MARGIN_INCHES,
)
from core.paper_structure import PAPER_OUTLINE, load_section
from core.bibliography import list_references, format_apa


def _create_base_document() -> Document:
    """Create a Document with standard academic formatting."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = DOCX_FONT
    font.size = Pt(DOCX_FONT_SIZE_PT)

    # Set paragraph spacing
    pf = style.paragraph_format
    pf.line_spacing = DOCX_LINE_SPACING
    pf.space_after = Pt(0)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(DOCX_MARGIN_INCHES)
        section.bottom_margin = Inches(DOCX_MARGIN_INCHES)
        section.left_margin = Inches(DOCX_MARGIN_INCHES)
        section.right_margin = Inches(DOCX_MARGIN_INCHES)

    return doc


def _add_markdown_to_doc(doc: Document, markdown_text: str) -> None:
    """Parse simple Markdown and add formatted content to the document."""
    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith("#### "):
            p = doc.add_heading(stripped[5:], level=4)
            _set_heading_font(p)
        elif stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
            _set_heading_font(p)
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
            _set_heading_font(p)
        elif stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
            _set_heading_font(p)
        # Bullet list
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
            _apply_inline_formatting(p)
        # Numbered list
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            p = doc.add_paragraph(text, style="List Number")
            _apply_inline_formatting(p)
        # Normal paragraph
        else:
            # Collect continuation lines
            para_lines = [stripped]
            while i + 1 < len(lines) and lines[i + 1].strip() and \
                    not lines[i + 1].strip().startswith("#") and \
                    not lines[i + 1].strip().startswith("- ") and \
                    not lines[i + 1].strip().startswith("* ") and \
                    not re.match(r"^\d+\.\s", lines[i + 1].strip()):
                i += 1
                para_lines.append(lines[i].strip())
            text = " ".join(para_lines)
            p = doc.add_paragraph()
            _add_formatted_runs(p, text)

        i += 1


def _set_heading_font(paragraph) -> None:
    """Ensure headings use the standard font."""
    for run in paragraph.runs:
        run.font.name = DOCX_FONT


def _apply_inline_formatting(paragraph) -> None:
    """Re-process paragraph text for bold/italic runs."""
    if paragraph.runs:
        text = paragraph.runs[0].text
        paragraph.clear()
        _add_formatted_runs(paragraph, text)


def _add_formatted_runs(paragraph, text: str) -> None:
    """Parse bold (**) and italic (*) Markdown and add as Word runs."""
    # Pattern matches **bold**, *italic*, or plain text
    pattern = r"(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))"
    for match in re.finditer(pattern, text):
        if match.group(2):  # Bold
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.font.name = DOCX_FONT
            run.font.size = Pt(DOCX_FONT_SIZE_PT)
        elif match.group(3):  # Italic
            run = paragraph.add_run(match.group(3))
            run.italic = True
            run.font.name = DOCX_FONT
            run.font.size = Pt(DOCX_FONT_SIZE_PT)
        elif match.group(4):  # Plain
            run = paragraph.add_run(match.group(4))
            run.font.name = DOCX_FONT
            run.font.size = Pt(DOCX_FONT_SIZE_PT)


def _add_references_chapter(doc: Document) -> None:
    """Append a References chapter with APA-formatted entries and hanging indent.

    Each entry uses a hanging indent: first line flush left, subsequent lines
    indented 0.5 inch.  The section is skipped entirely when the bibliography
    is empty.
    """
    refs = list_references()  # already sorted by citation key
    if not refs:
        return

    heading = doc.add_heading("References", level=1)
    _set_heading_font(heading)

    for ref in refs:
        apa_text = format_apa(ref)
        # Strip Markdown italic markers (*…*) — plain text only in docx
        apa_text = re.sub(r"\*(.+?)\*", r"\1", apa_text)

        p = doc.add_paragraph()
        pf = p.paragraph_format
        # Hanging indent: left indent 0.5 in, first-line indent -0.5 in
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.5)

        run = p.add_run(apa_text)
        run.font.name = DOCX_FONT
        run.font.size = Pt(DOCX_FONT_SIZE_PT)


def export_section(section_id: str, output_path: Path) -> Path:
    """Export a single section to .docx."""
    section = load_section(section_id)
    if not section:
        raise ValueError(f"Section '{section_id}' has no draft.")

    doc = _create_base_document()
    doc.add_heading(section["title"], level=1)
    _set_heading_font(doc.paragraphs[-1])
    _add_markdown_to_doc(doc, section["text"])

    output_path = Path(output_path)
    doc.save(str(output_path))
    return output_path


def export_full_paper(output_path: Path) -> Path:
    """Assemble all drafted sections into a single .docx."""
    doc = _create_base_document()

    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    from config import PAPER_TITLE
    run = title_para.add_run(PAPER_TITLE)
    run.bold = True
    run.font.name = DOCX_FONT
    run.font.size = Pt(16)
    doc.add_page_break()

    sections_written = 0
    first_section = True
    for section_id, title, _ in PAPER_OUTLINE:
        section = load_section(section_id)
        if not section:
            continue

        # Add section heading
        is_top_level = "." not in section_id
        level = 1 if is_top_level else 2

        # Page break before each top-level chapter, except the very first
        if is_top_level and not first_section:
            doc.add_page_break()

        doc.add_heading(title, level=level)
        _set_heading_font(doc.paragraphs[-1])
        _add_markdown_to_doc(doc, section["text"])
        sections_written += 1
        first_section = False

    if sections_written == 0:
        doc.add_paragraph("No sections have been drafted yet.")

    _add_references_chapter(doc)

    output_path = Path(output_path)
    doc.save(str(output_path))
    return output_path
